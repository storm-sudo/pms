from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('🚀 Synapse PMS - Deployment & Feature Summary', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Intro
    doc.add_paragraph('This document outlines the recent features added to the Synapse Project Management System and provides step-by-step instructions for deploying the Next.js application on Render.')

    # Phase 1
    doc.add_heading('🎨 Phase 1: Authentication (Email/Password)', level=1)
    doc.add_paragraph('We introduced a custom credentials system to securely manage access and conform to domain restrictions.')
    doc.add_paragraph('• Login & Registration: Supports custom credentials and session persistence.', style='List Bullet')
    doc.add_paragraph('• Safety Domain Policy: Enforces email validation strictly ending with correct handles (*nt@gmail.com).', style='List Bullet')
    doc.add_paragraph('• Sign Out Support: Quick sidebar option to securely log out.', style='List Bullet')

    # Phase 2
    doc.add_heading('📅 Phase 2: Overdue Automated Reminders (Resend API)', level=1)
    doc.add_paragraph('We added automated task due date and overdue email notifications for members and supervisors.')
    doc.add_paragraph('• Endpoint (/api/notify): Server-side handler utilizing Resend to send HTML alerts.', style='List Bullet')
    doc.add_paragraph('• Background Monitor: Scans all active tasks daily from dashboard layouts with daily rate-limiting to prevent inbox spams.', style='List Bullet')

    # Phase 3
    doc.add_heading('🔒 Phase 3: Access Control Validation Rules (RBAC)', level=1)
    doc.add_paragraph('We reinforced guards to ensure only authorized users can complete items.')

    # Table for RBAC
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component Action'
    hdr_cells[1].text = 'Restrictions / Allowed Users'

    row_data = [
        ('Closing Projects', 'Admins Only (YO, AR, SK)'),
        ('Closing Tasks', 'Admins Only (YO, AR, SK)'),
        ('Closing Subtasks', 'Supervisors Only (Lead or Assignee)')
    ]

    for item, restriction in row_data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = restriction

    # Configurations
    doc.add_heading('📝 Configuration Settings (.env.local)', level=1)
    doc.add_paragraph('Create an environment file to store secrets safely:')
    p = doc.add_paragraph()
    p.add_run('RESEND_API_KEY=your_resend_api_key_here').bold = True

    # Deployment
    doc.add_heading('🌐 Deployment Guidelines (Render)', level=1)
    doc.add_paragraph('Follow these steps to deploy on Render:')
    
    doc.add_paragraph('1. Create a New Web Service', style='List Number')
    doc.add_paragraph('Go to Render -> click New -> Web Service. Connect repo https://github.com/storm-sudo/pms.', style='List Bullet')
    
    doc.add_paragraph('2. Configure Service Defaults', style='List Number')
    
    # Render Settings Table
    table_render = doc.add_table(rows=1, cols=2)
    table_render.style = 'Table Grid'
    hdr_cells_r = table_render.rows[0].cells
    hdr_cells_r[0].text = 'Attribute'
    hdr_cells_r[1].text = 'Value'
    
    render_data = [
        ('Runtime', 'Node'),
        ('Build Command', 'npm run build'),
        ('Start Command', 'npm run start')
    ]
    for k, v in render_data:
        r_cells = table_render.add_row().cells
        r_cells[0].text = k
        r_cells[1].text = v

    doc.add_paragraph('3. Add Environment Variables', style='List Number')
    doc.add_paragraph('In Render\'s Environment tab, add RESEND_API_KEY with your key.', style='List Bullet')

    # Save
    filename = 'DEPLOYMENT_SUMMARY.docx'
    doc.save(filename)
    print(f"[{filename}] successfully created!")

if __name__ == '__main__':
    create_document()
